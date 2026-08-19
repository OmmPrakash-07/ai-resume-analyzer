import re
from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def clean_text(text: str) -> str:
    """Remove unnecessary spaces while preserving separate lines."""

    text = text.replace("\xa0", " ")
    text = text.replace("\u200b", "")

    cleaned_lines = []

    for line in text.splitlines():
        line = re.sub(r"[ \t]+", " ", line).strip()

        if line:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract readable text from a PDF file."""

    pages = []

    with pymupdf.open(stream=file_bytes, filetype="pdf") as document:
        if document.needs_pass:
            raise ValueError("The PDF is password protected.")

        for page in document:
            page_text = page.get_text("text", sort=True)

            if page_text.strip():
                pages.append(page_text)

    extracted_text = clean_text("\n".join(pages))

    if not extracted_text:
        raise ValueError(
            "No readable text was found. The PDF may be scanned or image-based."
        )

    return extracted_text


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract paragraphs, tables, headers and footers from DOCX."""

    document = Document(BytesIO(file_bytes))
    extracted_parts = []

    # Normal paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            extracted_parts.append(paragraph.text)

    # Text inside tables
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )

            if row_text:
                extracted_parts.append(row_text)

    # Headers and footers
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                extracted_parts.append(paragraph.text)

        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                extracted_parts.append(paragraph.text)

    extracted_text = clean_text("\n".join(extracted_parts))

    if not extracted_text:
        raise ValueError("No readable text was found in the DOCX file.")

    return extracted_text


def extract_resume_text(uploaded_file) -> dict:
    """Select the appropriate extractor based on the uploaded file."""

    filename = uploaded_file.name
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Only PDF and DOCX files are supported.")

    file_bytes = uploaded_file.getvalue()

    if not file_bytes:
        raise ValueError("The uploaded file is empty.")

    if extension == ".pdf":
        text = extract_pdf_text(file_bytes)
    else:
        text = extract_docx_text(file_bytes)

    return {
        "filename": filename,
        "extension": extension,
        "text": text,
        "characters": len(text),
        "words": len(text.split()),
    }